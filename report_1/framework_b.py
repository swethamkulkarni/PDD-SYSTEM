"""
framework_b.py
Framework B — 8-section VC IC Memo builder (skeptical senior partner tone).

Each narrative section is written by the LLM in the voice of a skeptical
senior VC partner. Structured pipeline data (findings tables, risk snapshots)
remains rule-based.

Usage:
    from framework_b import build_framework_b
    doc = build_framework_b(result)
    doc.save("ic_memo.docx")

    # Explicit LLM provider:
    doc = build_framework_b(result, provider="groq")

Dependencies:
    pip install python-docx
"""

from __future__ import annotations

from docx import Document
import os

# ADDED: 24/05
from report_1.doc_helpers import (
    configure_styles, cover_page, title_block, extract_company_name,
    setup_page_header, setup_page_footer,
    add_heading, add_body, add_bullet, add_labelled_bullet, add_placeholder,
    add_mixed_section, add_gap_note, add_questions_list,
    add_findings_table, add_questions_table,
    add_card, add_severity_header,
    safe_call,
    SEVERITY_COLOUR, SEVERITY_LABEL, RED, GREEN, DARK_GREY, AMBER,
)

from report_1.section_writer import (
    SectionWriter,
    collect_slide_texts,
    collect_unknown_texts,
    collect_slide_tables,
    findings_summary_for,
    all_findings_summary,
)


def _write(writer: SectionWriter, doc: Document, section_key: str,
           slide_texts: list, unknown_texts: list,
           findings_summary: str = "", framework: str = "B",
           max_tokens: int = 600, slide_tables: list = None) -> None:
    """Generate LLM prose (IC memo tone) and write using the mixed (bullet-aware) renderer."""
    text = writer.write_section(
        section_key, slide_texts, unknown_texts,
        findings_summary, framework, max_tokens,
        slide_tables=slide_tables,
    )
    add_mixed_section(doc, text)


def _gap(doc: Document, message: str) -> None:
    """Shorthand for adding a material information gap note."""
    add_gap_note(doc, message)


def build_framework_b(result, provider: str = None) -> Document:
    """
    Build the 8-section IC Memo and return a python-docx Document.

    Written in the tone of a skeptical senior VC partner.

    Sections
    --------
    1. Executive Summary
    2. Market & Problem
    3. Technology & Product
    4. Financials
    5. GTM & Traction
    6. Competitive Landscape
    7. Comments & Observations
    8. IC Memo Summary  (positive signals / key risks / recommendation)
    """
    doc = Document()
    configure_styles(doc)

    source       = getattr(getattr(result, "deck", object()), "source_filename", "Unknown")
    all_findings = getattr(result, "all_findings", []) or []
    flagged      = safe_call(result, "flagged")            or []
    unclear      = safe_call(result, "unclear")            or []
    ic_risks     = safe_call(result, "ic_memo_risks")      or []
    questions    = safe_call(result, "questions_by_category") or {}
    high_findings = [f for f in all_findings if getattr(f, "severity", "") == "HIGH"]
    med_findings  = [f for f in all_findings if getattr(f, "severity", "") == "MEDIUM"]

    #EDIT:10/04
    doc          = Document()
    LOGO_PATH    = os.path.join(os.path.dirname(__file__), "assets", "logo.png")
    HERO_PATH    = os.path.join(os.path.dirname(__file__), "assets", "cover_photo.jpg")

    configure_styles(doc)

    section = doc.sections[0]
    section.different_first_page_header_footer = True
    setup_page_header(section, logo_image_path=LOGO_PATH if os.path.isfile(LOGO_PATH) else None)
    setup_page_footer(section)

# EDIT: 04/06
    source       = getattr(getattr(result, "deck", object()), "source_filename", "Unknown")
    writer       = SectionWriter(provider=provider)
    company_name = writer.get_company_name(result)

# ADDED: 24/05
    title_block(
        doc,
        company_name    = company_name,
        report_type     = "Investment Committee Memo  |  Confidential — Internal Use Only",
        source_filename = source,
    )
    

    writer   = SectionWriter(provider=provider)
    unknowns = collect_unknown_texts(result)

    # Helper: gather texts across multiple section keys
    def _slides(*keys):
        texts = []
        for k in keys:
            texts += collect_slide_texts(result, k)
        return texts

    def _tables(*keys):
        tbls = []
        for k in keys:
            tbls += collect_slide_tables(result, k)
        return tbls

    # ── 1. Executive Summary ────────────────────────────────────────────────
    add_heading(doc, "1. Executive Summary", level=1)
    exec_slides = _slides("cover", "market", "solution", "problem", "product", "team")
    _write(writer, doc, "executive_summary_b",
           slide_texts=exec_slides or unknowns[:4],
           unknown_texts=unknowns,
           findings_summary=all_findings_summary(result),
           max_tokens=600)

    add_heading(doc, "Headline Risk Snapshot", level=2)
    add_body(doc,
             f"Total checks: {len(all_findings)}   |   Flagged: {len(flagged)}   |   "
             f"Unclear: {len(unclear)}   |   HIGH: {len(high_findings)}   |   "
             f"MEDIUM: {len(med_findings)}")
    if high_findings:
        add_body(doc, "HIGH severity issues requiring IC attention:", bold=True, colour=RED)
        for f in high_findings[:6]:
            add_bullet(
                doc,
                f"{getattr(f, 'label', '')} — {(getattr(f, 'evidence', '') or '')[:120]}",
                colour=RED,
            )

    # ── 2. Market & Problem ──────────────────────────────────────────────────
    add_heading(doc, "2. Market & Problem", level=1)
    mkt_slides = _slides("market", "market opportunity", "problem", "tam",
                         "industry", "landscape", "cover")
    _write(writer, doc, "market_opportunity",
           slide_texts=mkt_slides or unknowns[:3],
           unknown_texts=unknowns,
           findings_summary=findings_summary_for(result, "market", "ms", "tam"),
           max_tokens=600)
    if not mkt_slides:
        _gap(doc, "No market opportunity or problem slides identified — market sizing unverified.")

    market_flags = [
        f for f in all_findings
        if "market" in (getattr(f, "category", "") or "").lower()
        and getattr(f, "severity", "") in ("HIGH", "MEDIUM")
    ]
    if market_flags:
        add_heading(doc, "Market-Related Concerns", level=2)
        for f in market_flags:
            sev = getattr(f, "severity", "")
            add_labelled_bullet(doc, SEVERITY_LABEL.get(sev, sev),
                                f" {getattr(f, 'label', '')}",
                                colour=SEVERITY_COLOUR.get(sev, DARK_GREY))

    # ── 3. Technology & Product ──────────────────────────────────────────────
    add_heading(doc, "3. Technology & Product", level=1)
    tech_slides = _slides("solution", "technology", "product", "platform",
                          "architecture", "ai", "ml")
    _write(writer, doc, "technology_solution",
           slide_texts=tech_slides or unknowns[:3],
           unknown_texts=unknowns,
           findings_summary=findings_summary_for(result, "tech", "ip", "ms", "product"),
           max_tokens=600)
    if not tech_slides:
        _gap(doc, "Technical architecture, AI/ML detail, and IP defensibility not described in the deck.")

    tech_flags = [
        f for f in all_findings
        if any(k in (getattr(f, "category", "") or "").lower()
               for k in ("tech", "product", "ip"))
        and getattr(f, "severity", "") in ("HIGH", "MEDIUM")
    ]
    if tech_flags:
        add_heading(doc, "Technology Risk Flags", level=2)
        for f in tech_flags:
            sev = getattr(f, "severity", "")
            add_labelled_bullet(doc, SEVERITY_LABEL.get(sev, sev),
                                f" {getattr(f, 'label', '')} — "
                                f"{(getattr(f, 'evidence', '') or '')[:120]}",
                                colour=SEVERITY_COLOUR.get(sev, DARK_GREY))

    # ── 4. Financials ────────────────────────────────────────────────────────
    add_heading(doc, "4. Financials", level=1)
    add_body(doc,
             "All financial data is sourced from the pitch deck and has not been "
             "independently audited. The committee should request supporting schedules "
             "before placing material reliance on these figures.",
             italic=True)
    fin_slides = _slides("financials", "cap_table", "traction", "revenue",
                         "funding", "use of proceeds", "investment", "raise")
    _write(writer, doc, "comments_financials",
           slide_texts=fin_slides or unknowns[:3],
           unknown_texts=unknowns,
           findings_summary=findings_summary_for(result, "financ", "revenue", "valuat", "cap", "fund"),
           slide_tables=_tables("financials", "cap_table", "traction"),
           max_tokens=600)
    if not fin_slides:
        _gap(doc, "No financial slides identified — revenue model, burn rate, and runway unknown.")

    fin_flags = [
        f for f in all_findings
        if any(k in (getattr(f, "category", "") or "").lower()
               for k in ("financ", "revenue", "valuat", "cap", "funding"))
    ]
    if fin_flags:
        add_heading(doc, "Financial Anomaly Findings", level=2)
        add_findings_table(doc, fin_flags)

    # ── 5. GTM & Traction ────────────────────────────────────────────────────
    add_heading(doc, "5. GTM & Traction", level=1)
    tract_slides = _slides("traction", "financials", "go-to-market", "gtm",
                           "customers", "sales", "growth", "metrics", "kpi")
    _write(writer, doc, "traction",
           slide_texts=tract_slides or unknowns[:3],
           unknown_texts=unknowns,
           findings_summary=findings_summary_for(result, "traction", "gtm", "customer", "metric", "growth"),
           max_tokens=600)
    if not tract_slides:
        _gap(doc, "No traction or GTM slides identified — sales pipeline and commercial evidence absent.")

    gtm_flags = [
        f for f in all_findings
        if any(k in (getattr(f, "category", "") or "").lower()
               for k in ("traction", "gtm", "customer", "metric", "growth"))
    ]
    if gtm_flags:
        add_heading(doc, "GTM / Traction Concerns", level=2)
        for f in gtm_flags:
            sev = getattr(f, "severity", "")
            add_labelled_bullet(doc, SEVERITY_LABEL.get(sev, sev),
                                f" {getattr(f, 'label', '')}",
                                colour=SEVERITY_COLOUR.get(sev, DARK_GREY))

# ADDED: 24/05
    # DD questions rendered as severity-grouped cards
    critical_qs = [f for f in all_findings
                   if getattr(f, "severity", "") == "HIGH"
                   and getattr(f, "ic_memo_flag", False)]
    high_qs     = [f for f in all_findings
                   if getattr(f, "severity", "") == "HIGH"
                   and not getattr(f, "ic_memo_flag", False)]
    medium_qs   = [f for f in all_findings
                   if getattr(f, "severity", "") == "MEDIUM"]

    if critical_qs or high_qs or medium_qs:
        add_heading(doc, "Key Diligence Questions", level=2)

    if critical_qs:
        add_severity_header(doc, "CRITICAL", len(critical_qs),
                            "Potential deal-breakers. Must be resolved before any capital commitment.")
        for i, f in enumerate(critical_qs, 1):
            add_card(doc, i,
                     category  = getattr(f, "category", ""),
                     source    = f"Slide {getattr(f, 'slide_reference', 'N/A')}",
                     body_text = getattr(f, "generated_question", "") or getattr(f, "label", ""),
                     severity  = "CRITICAL")

    if high_qs:
        add_severity_header(doc, "HIGH", len(high_qs),
                            "Material to investment decision. Must resolve before close.")
        for i, f in enumerate(high_qs, 1):
            add_card(doc, i,
                     category  = getattr(f, "category", ""),
                     source    = f"Slide {getattr(f, 'slide_reference', 'N/A')}",
                     body_text = getattr(f, "generated_question", "") or getattr(f, "label", ""),
                     severity  = "HIGH")

    if medium_qs:
        add_severity_header(doc, "MEDIUM", len(medium_qs),
                            "Important for post-close monitoring and ongoing investor reporting.")
        for i, f in enumerate(medium_qs, 1):
            add_card(doc, i,
                     category  = getattr(f, "category", ""),
                     source    = f"Slide {getattr(f, 'slide_reference', 'N/A')}",
                     body_text = getattr(f, "generated_question", "") or getattr(f, "label", ""),
                     severity  = "MEDIUM")

    # ── 6. Competitive Landscape ─────────────────────────────────────────────
    add_heading(doc, "6. Competitive Landscape", level=1)
    comp_slides = _slides("market", "technology", "competition", "competitive",
                          "landscape", "differentiation", "solution")
    _write(writer, doc, "competitive_landscape",
           slide_texts=comp_slides or unknowns[:3],
           unknown_texts=unknowns,
           findings_summary=findings_summary_for(result, "competi", "ms", "market"),
           max_tokens=600)
    if not comp_slides:
        _gap(doc, "No competitive landscape slide identified — differentiation claims unverified.")

    comp_flags = [
        f for f in all_findings
        if "competi" in (getattr(f, "category", "") or "").lower()
    ]
    if comp_flags:
        add_heading(doc, "Competitive Risk Flags", level=2)
        for f in comp_flags:
            sev = getattr(f, "severity", "")
            add_labelled_bullet(doc, SEVERITY_LABEL.get(sev, sev),
                                f" {getattr(f, 'label', '')}",
                                colour=SEVERITY_COLOUR.get(sev, DARK_GREY))

    # ── 7. Comments & Observations ───────────────────────────────────────────
    add_heading(doc, "7. Comments & Observations", level=1)

    obs_map = [
        ("7.1 Technology",  "comments_technology",
         ["technology", "solution", "platform", "product", "architecture"],
         ["tech", "ip", "product"]),
        ("7.2 Financials",  "comments_financials",
         ["financials", "cap_table", "traction", "revenue", "funding"],
         ["financ", "revenue", "valuat", "cap"]),
        ("7.3 Regulations", "comments_regulations",
         ["market", "regulatory", "compliance", "legal"],
         ["regulat", "legal", "compliance"]),
        ("7.4 GTM",         "comments_gtm",
         ["traction", "market", "go-to-market", "gtm", "customers", "sales"],
         ["traction", "gtm", "customer", "sales"]),
        ("7.5 Competition", "comments_competition",
         ["market", "technology", "competition", "competitive", "differentiation"],
         ["competi", "ms", "market"]),
        ("7.6 Other",       "areas_to_watch",
         ["cover"],                       []),
    ]

    for heading_text, key, slide_keys, finding_keys in obs_map:
        add_heading(doc, heading_text, level=2)
        slides = _slides(*slide_keys)
        _write(writer, doc, key,
               slide_texts=slides or unknowns[:2],
               unknown_texts=unknowns,
               findings_summary=findings_summary_for(result, *finding_keys) if finding_keys
                                else all_findings_summary(result),
               max_tokens=450)

    # ── 8. IC Memo Summary ───────────────────────────────────────────────────
    add_heading(doc, "8. IC Memo Summary", level=1)

    # 8.1 Positive Signals
    add_heading(doc, "8.1 Positive Signals", level=2)
    positive_lines = writer.write_lines(
        "ic_summary_positive",
        slide_texts=_slides("cover", "market", "solution", "team",
                            "traction", "technology"),
        unknown_texts=unknowns,
        findings_summary=all_findings_summary(result),
        framework="B",
        max_tokens=400,
    )
    for line in positive_lines:
        content = line.lstrip("- ").strip()
        if content:
            add_bullet(doc, content, colour=GREEN)

    # 8.2 Key Risks
    add_heading(doc, "8.2 Key Risks", level=2)

    key_risks = list(ic_risks) + [f for f in high_findings if f not in ic_risks]
    if key_risks:
        for f in key_risks[:10]:
            sev = getattr(f, "severity", "")
            add_labelled_bullet(doc, SEVERITY_LABEL.get(sev, sev),
                                f" {getattr(f, 'label', '')} — "
                                f"{(getattr(f, 'evidence', '') or '')[:120]}",
                                colour=SEVERITY_COLOUR.get(sev, DARK_GREY))

    risk_lines = writer.write_lines(
        "ic_summary_risks",
        slide_texts=_slides("cover", "financials", "traction", "market",
                            "technology", "competition"),
        unknown_texts=unknowns,
        findings_summary=all_findings_summary(result),
        framework="B",
        max_tokens=400,
    )
    for line in risk_lines:
        content = line.lstrip("- ").strip()
        if content:
            add_bullet(doc, content, colour=RED)

    # 8.3 Recommendation
    add_heading(doc, "8.3 Recommendation", level=2)
    rec = writer.write_section(
        "ic_recommendation",
        slide_texts=_slides("cover", "market", "financials", "traction",
                            "solution", "team", "technology"),
        unknown_texts=unknowns,
        findings_summary=all_findings_summary(result),
        framework="B",
        max_tokens=150,
    )
    add_body(doc, rec, bold=True)

    return doc