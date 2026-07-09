"""
framework_a.py
Framework A — 11-section Due Diligence Report builder.

Each narrative section is written by the LLM (via SectionWriter).
Structured data (findings tables, questions, risk lists) remains rule-based.

Usage:
    from framework_a import build_framework_a
    doc = build_framework_a(result)
    doc.save("dd_report.docx")

    # Explicit LLM provider:
    doc = build_framework_a(result, provider="groq")

Dependencies:
    pip install python-docx
"""

from __future__ import annotations

from docx import Document

#EDIT: 10/04: moved all doc-building helper functions and constants to doc_helpers.py for better organization and reuse across different frameworks.
import os

# ADDED: 24/05
from report_1.doc_helpers import (
    configure_styles, cover_page, title_block, extract_company_name,
    setup_page_header, setup_page_footer,
    add_heading, add_body, add_bullet, add_labelled_bullet, add_placeholder,
    add_mixed_section, add_gap_note, add_questions_list,
    add_findings_table, add_questions_table,
    add_card, add_severity_header,
    safe_call,
    SEVERITY_COLOUR, SEVERITY_LABEL, RED, GREEN, DARK_GREY, AMBER,
)


from report_1.section_writer import (
    SectionWriter,
    collect_slide_texts,
    collect_unknown_texts,
    collect_slide_tables,
    findings_summary_for,
    all_findings_summary,
)

def _write(writer: SectionWriter, doc: Document, section_key: str,
           slide_texts: list, unknown_texts: list,
           findings_summary: str = "", framework: str = "A",
           max_tokens: int = 700, slide_tables: list = None,
           on_complete=None) -> None:
    """Generate LLM prose and write it using the mixed (bullet-aware) renderer."""
    text = writer.write_section(
        section_key, slide_texts, unknown_texts,
        findings_summary, framework, max_tokens,
        slide_tables=slide_tables,
    )
    add_mixed_section(doc, text)
    if on_complete:
        on_complete(section_key)


def _gap(doc: Document, message: str) -> None:
    """Shorthand for adding a material information gap note."""
    add_gap_note(doc, message)

#edit: added _write_questions for the "Investor Questions" section, which uses a different renderer that formats the output as a numbered list instead of mixed prose and bullets.
def _write_questions(writer: SectionWriter, doc: Document, section_key: str,
                     slide_texts: list, unknown_texts: list,
                     findings_summary: str = "", max_tokens: int = 500) -> None:
    """Generate a numbered question list and render it with add_questions_list."""
    text = writer.write_section(
        section_key, slide_texts, unknown_texts,
        findings_summary, "A", max_tokens,
    )
    add_questions_list(doc, text)

def build_framework_a(result, provider: str = None, on_section_complete=None) -> Document:
    """
    Build the 11-section Due Diligence Report and return a python-docx Document.

    Sections
    --------
    1.  Product Description
    2.  Technology Solution
    3.  Business Model
    4.  Technology Overview
    5.  Team
    6.  Cap Table & Funding
    7.  Traction
    8.  Investment Rationale
    9.  Areas to Watch          (LLM narrative + rule-based findings table)
    10. Comments & Observations (LLM per sub-section)
    11. Final Call              (LLM recommendation)
    """
    #EDIT:10/04
    doc = Document()

    # EDIT:10/04 Optional: set these paths to your actual asset files
    LOGO_PATH    = os.path.join(os.path.dirname(__file__), "assets", "logo.png")
    HERO_PATH    = os.path.join(os.path.dirname(__file__), "assets", "cover_photo.jpg")

    configure_styles(doc)

    # Wire inner-page header (logo top-right) and footer (contacts + copyright).
    # different_first_page_header_footer = True keeps the cover page clean.
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    setup_page_header(section, logo_image_path=LOGO_PATH if os.path.isfile(LOGO_PATH) else None)
    setup_page_footer(section)

#EDIT: 04/06
    source       = getattr(getattr(result, "deck", object()), "source_filename", "Unknown")
    writer       = SectionWriter(provider=provider)
    company_name = writer.get_company_name(result)
    

# EDIT: 24/05
    title_block(
        doc,
        company_name  = company_name,
        report_type   = "Preliminary Due Diligence Report",
        source_filename = source,
    )

    writer   = SectionWriter(provider=provider)
    unknowns = collect_unknown_texts(result)

    all_findings = getattr(result, "all_findings", []) or []
    flagged      = safe_call(result, "flagged")        or []
    unclear      = safe_call(result, "unclear")        or []
    ic_risks     = safe_call(result, "ic_memo_risks")  or []
    questions    = safe_call(result, "questions_by_category") or {}
    total_high   = len([f for f in all_findings if getattr(f, "severity", "") == "HIGH"])
    total_med    = len([f for f in all_findings if getattr(f, "severity", "") == "MEDIUM"])

    # Helper: gather texts across multiple section keys
    def _slides(*keys):
        texts = []
        for k in keys:
            texts += collect_slide_texts(result, k)
        return texts

    def _tables(*keys):
        tbls = []
        for k in keys:
            tbls += collect_slide_tables(result, k)
        return tbls

    # ── 1. Product Description ──────────────────────────────────────────────
    add_heading(doc, "1. Product Description", level=1)
    prod_slides = _slides("cover", "problem", "solution", "market", "product")
    _write(writer, doc, "product_description",
           slide_texts=prod_slides or unknowns[:3],
           unknown_texts=unknowns,
           findings_summary=findings_summary_for(result, "ms", "product"),
           max_tokens=700, on_complete=on_section_complete)
    if not prod_slides:
        _gap(doc, "No dedicated product or problem slides identified — review UNKNOWN slides.")

    # ── 2. Technology Solution ──────────────────────────────────────────────
    add_heading(doc, "2. Technology Solution", level=1)
    tech_slides = _slides("solution", "technology", "product", "platform")
    _write(writer, doc, "technology_solution",
           slide_texts=tech_slides or unknowns[:3],
           unknown_texts=unknowns,
           findings_summary=findings_summary_for(result, "tech", "ip", "product"),
           max_tokens=700, on_complete=on_section_complete)
    if not tech_slides:
        _gap(doc, "Technical architecture, AI/ML specifics, and IP position not described in deck.")

    # ── 3. Business Model ───────────────────────────────────────────────────
    add_heading(doc, "3. Business Model", level=1)
    biz_slides = _slides("financials", "business model", "revenue", "traction",
                         "go-to-market", "gtm", "pricing")
    _write(writer, doc, "business_model",
           slide_texts=biz_slides or unknowns[:3],
           unknown_texts=unknowns,
           findings_summary=findings_summary_for(result, "financ", "revenue", "business", "pricing"),
           max_tokens=700, on_complete=on_section_complete)
    if not biz_slides:
        _gap(doc, "Revenue streams, pricing model, and unit economics not explicitly presented.")

    # ── 4. Technology Overview ──────────────────────────────────────────────
    add_heading(doc, "4. Technology Overview", level=1)

    add_heading(doc, "4.1 Technical Architecture", level=2)
    arch_slides = _slides("technology", "solution", "product", "platform", "architecture")
    _write(writer, doc, "technology_overview",
           slide_texts=arch_slides or unknowns[:3],
           unknown_texts=unknowns,
           findings_summary=findings_summary_for(result, "tech", "ip", "product"),
           max_tokens=700, on_complete=on_section_complete)

    add_heading(doc, "4.2 IP & Defensibility", level=2)
    _write(writer, doc, "comments_technology",
           slide_texts=_slides("technology", "solution", "ip"),
           unknown_texts=unknowns,
           findings_summary=findings_summary_for(result, "ip", "competi", "tech"),
           max_tokens=500, on_complete=on_section_complete)
    if not arch_slides:
        _gap(doc, "Scalability, AI/ML detail, and IP defensibility not addressed in the deck.")

    # ── 5. Team ─────────────────────────────────────────────────────────────
    add_heading(doc, "5. Team", level=1)
    team_slides = _slides("team", "cover", "founders", "about", "management")
    _write(writer, doc, "team",
           slide_texts=team_slides or unknowns[:3],
           unknown_texts=unknowns,
           findings_summary=findings_summary_for(result, "team", "ms", "founder"),
           max_tokens=600, on_complete=on_section_complete)
    if not team_slides:
        _gap(doc, "No dedicated team slide identified — team composition and background unclear.")

    # ── 6. Cap Table & Funding ──────────────────────────────────────────────
    add_heading(doc, "6. Cap Table & Funding", level=1)
    cap_slides = _slides("cap_table", "funding", "financials", "traction",
                         "investment", "use of proceeds", "raise")
    _write(writer, doc, "cap_table",
           slide_texts=cap_slides or unknowns[:3],
           unknown_texts=unknowns,
           findings_summary=findings_summary_for(result, "cap", "fund", "financ", "valuat"),
           slide_tables=_tables("cap_table", "funding", "financials"),
           max_tokens=600, on_complete=on_section_complete)
    if not cap_slides:
        _gap(doc, "Cap table, ownership structure, and valuation not presented in the deck — material gap.")

    # ── 7. Traction ─────────────────────────────────────────────────────────
    add_heading(doc, "7. Traction", level=1)

    add_heading(doc, "7.1 Key Metrics", level=2)
    tract_slides = _slides("traction", "metrics", "financials", "customers",
                           "growth", "revenue", "kpi")
    _write(writer, doc, "traction",
           slide_texts=tract_slides or unknowns[:3],
           unknown_texts=unknowns,
           findings_summary=findings_summary_for(result, "traction", "metric", "revenue", "customer"),
           max_tokens=600, on_complete=on_section_complete)

    add_heading(doc, "7.2 Customer & Revenue Growth", level=2)
    _write(writer, doc, "business_model",
           slide_texts=_slides("financials", "traction", "customers", "revenue"),
           unknown_texts=unknowns,
           findings_summary=findings_summary_for(result, "revenue", "customer", "growth"),
           max_tokens=500, on_complete=on_section_complete)
    if not tract_slides:
        _gap(doc, "No concrete traction metrics (ARR, MRR, customers) presented — significant gap for DD.")

    # ── 8. Investment Rationale ──────────────────────────────────────────────
    add_heading(doc, "8. Investment Rationale", level=1)

    add_heading(doc, "8.1 Market Opportunity", level=2)
    mkt_slides = _slides("market", "market opportunity", "tam", "problem",
                         "industry", "landscape")
    _write(writer, doc, "market_opportunity",
           slide_texts=mkt_slides or unknowns[:3],
           unknown_texts=unknowns,
           findings_summary=findings_summary_for(result, "market", "tam"),
           max_tokens=600, on_complete=on_section_complete)

    add_heading(doc, "8.2 Competitive Landscape", level=2)
    comp_slides = _slides("market", "competition", "competitive", "landscape",
                          "technology", "differentiation")
    _write(writer, doc, "competitive_landscape",
           slide_texts=comp_slides or unknowns[:3],
           unknown_texts=unknowns,
           findings_summary=findings_summary_for(result, "competi", "ms", "market"),
           max_tokens=600, on_complete=on_section_complete)
    if not comp_slides:
        _gap(doc, "No competitive analysis slide identified — differentiation claims unverified.")

    add_heading(doc, "8.3 Investment Rationale", level=2)
    _write(writer, doc, "investment_rationale",
           slide_texts=_slides("cover", "market", "solution", "traction",
                               "team", "investment", "thesis"),
           unknown_texts=unknowns,
           findings_summary=all_findings_summary(result),
           max_tokens=600, on_complete=on_section_complete)

#edit: expanded section 9 to cover multiple key risk areas (technology, financials, regulations, GTM, competition) with dedicated LLM-generated narratives for each, in addition to the overall "Areas to Watch" section. Each sub-section pulls from relevant slide texts and findings to provide a focused analysis of risks and concerns in that area.
# ── 9. Areas to Watch ───────────────────────────────────────────────────
    add_heading(doc, "9. Areas to Watch", level=1)
    _write(writer, doc, "areas_to_watch",
           slide_texts=_slides("cover", "market", "financials", "traction",
                               "risk", "solution", "technology"),
           unknown_texts=unknowns,
           findings_summary=all_findings_summary(result),
           max_tokens=700, on_complete=on_section_complete)
    
    # ADDED: 24/05: Render flagged findings as cards
    if flagged or unclear:
        add_heading(doc, "Pipeline Risk Flags", level=2)
        sev_map = {"HIGH": "CRITICAL", "MEDIUM": "HIGH", "LOW": "MEDIUM"}
        for i, f in enumerate((flagged + unclear)[:20], 1):
            add_card(
                doc,
                number    = i,
                category  = getattr(f, "category", ""),
                source    = f"Slide {getattr(f, 'slide_reference', 'N/A')}",
                body_text = f"{getattr(f, 'label', '')} — {(getattr(f, 'evidence', '') or '')[:300]}",
                severity  = sev_map.get(getattr(f, "severity", ""), "MEDIUM"),
            )

#edit: added section 10 with multiple sub-sections for detailed comments and observations across key areas, allowing the LLM to provide nuanced insights based on the slide content and findings in each domain. This structured approach helps ensure that the report captures a comprehensive view of the startup's strengths, weaknesses, and areas of concern from multiple angles.
# ── 10. Comments & Observations ─────────────────────────────────────────
    add_heading(doc, "10. Comments & Observations", level=1)

    obs_map = [
        ("10.1 Technology",   "comments_technology",
         ["technology", "solution", "platform", "product", "architecture"],
         ["tech", "ip", "product"]),
        ("10.2 Financials & Commercials", "comments_financials",
         ["financials", "cap_table", "traction", "revenue", "funding"],
         ["financ", "revenue", "valuat", "cap"]),
        ("10.3 Regulations",  "comments_regulations",
         ["market", "regulatory", "compliance", "legal"],
         ["regulat", "legal", "compliance"]),
        ("10.4 GTM Strategy", "comments_gtm",
         ["traction", "market", "go-to-market", "gtm", "customers", "sales"],
         ["traction", "gtm", "customer", "sales"]),
        ("10.5 Competition (Benchmarking)", "comments_competition",
         ["market", "technology", "competition", "competitive", "differentiation"],
         ["competi", "ms", "market"]),
    ]

    for heading_text, key, slide_keys, finding_keys in obs_map:
        add_heading(doc, heading_text, level=2)
        slides = _slides(*slide_keys)
        _write_questions(writer, doc, key,
                         slide_texts=slides or unknowns[:2],
                         unknown_texts=unknowns,
                         findings_summary=findings_summary_for(result, *finding_keys),
                         max_tokens=500)

    add_heading(doc, "10.6 Other Areas", level=2)
    _write_questions(writer, doc, "comments_technology",
                     slide_texts=_slides("cover") or unknowns[:2],
                     unknown_texts=unknowns,
                     findings_summary=all_findings_summary(result),
                     max_tokens=400)

    # ── 11. Final Call ───────────────────────────────────────────────────────
    add_heading(doc, "11. Final Call", level=1)
    
#edit: expanded the "Final Call" section to pull from a broader range of slide texts and findings, allowing the LLM to synthesize a more comprehensive recommendation based on the overall analysis. This approach helps ensure that the final investment recommendation is well-informed by all aspects of the due diligence process, including market opportunity, technology, team, traction, and identified risks.
    _write(writer, doc, "final_call_a",
           slide_texts=_slides("cover", "market", "financials", "traction",
                               "solution", "team", "technology"),
           unknown_texts=unknowns,
           findings_summary=all_findings_summary(result),
           max_tokens=700)


    return doc