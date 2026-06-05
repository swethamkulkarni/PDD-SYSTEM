"""
doc_helpers.py
Shared colours, severity maps, and Word-document utility functions used by
both Framework A and Framework B builders.

Dependencies:
    pip install python-docx
"""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.shared import OxmlElement as _OxmlElement
from datetime import datetime
import os


#edit: add _FALLBACK_NOTE_SENTINEL for consistent detection of LLM fallback responses in add_mixed_section() and add_questions_list().
# Sentinel string matched by rendering helpers to detect LLM fallback responses
_FALLBACK_NOTE_SENTINEL = "[LLM unavailable — review and complete this section manually.]"

# ---------------------------------------------------------------------------
# Colour palette
#edit:24/05
# ---------------------------------------------------------------------------
NAVY       = RGBColor(0x0F, 0x4C, 0x5C)   # dark teal — headings, card numbers
STEEL      = RGBColor(0x0F, 0x4C, 0x5C)   # same teal for sub-headings
DARK_GREY  = RGBColor(0x5C, 0x5C, 0x5C)   # body meta text
RED        = RGBColor(0xC0, 0x39, 0x2B)   # CRITICAL severity
AMBER      = RGBColor(0xE8, 0x7C, 0x1E)   # HIGH severity / gap warnings
GREEN      = RGBColor(0x27, 0xAE, 0x60)   # MEDIUM / pass
CYAN       = RGBColor(0x00, 0xB0, 0xF0)   # StartUpScale360 brand cyan (keep for cover)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)   # white text on dark backgrounds
BLACK      = RGBColor(0x00, 0x00, 0x00)   # company name title
LIGHT_GREY = RGBColor(0x70, 0x70, 0x70)   # dates / footer meta text



LIGHT_BLUE_FILL   = "E8F0F2"   # table header fill — matches sample teal-grey
ALT_ROW_FILL      = "F5F8F9"   # alternating row fill
CARD_NUMBER_FILL  = "0F4C5C"   # dark teal — card number left cell
CARD_HEADER_FILL  = "E8F0F2"   # light teal-grey — card header right cell
CRITICAL_FILL     = "C0392B"   # red — CRITICAL severity header
HIGH_FILL         = "E87C1E"   # amber — HIGH severity header
MEDIUM_FILL       = "27AE60"   # green — MEDIUM severity header

# Card column widths in DXA (A4 with 0.5" margins = 9638 DXA content width)
_CARD_NUM_COL  = 771    # ~8% — left number cell
_CARD_BODY_COL = 8867   # ~92% — right content cell

#EDIT:11/04
# ---------------------------------------------------------------------------
# StartUpScale360 contact details — edit these to update all headers/footers
# ---------------------------------------------------------------------------
COMPANY_NAME    = "StartUpScale360 FZE"
COMPANY_TEL     = "+971 XX XXX XXXX"          # ← replace with real number
COMPANY_ADDRESS = "IFZA Business Park, Dubai"  # ← replace with real address
COMPANY_WEBSITE = "www.startupscale360.com"    # ← replace with real URL
COMPANY_EMAIL   = "info@startupscale360.com"   # ← replace with real email
COPYRIGHT_LINE  = (
    "© 2025 StartUpScale360 FZE. This report is confidential and intended solely "
    "for the named recipient. Not to be reproduced or distributed without prior "
    "written consent."
)

# ---------------------------------------------------------------------------
# Severity helpers
# ---------------------------------------------------------------------------
SEVERITY_COLOUR = {
    "HIGH":   RED,
    "MEDIUM": AMBER,
    "LOW":    GREEN,
    "PASS":   GREEN,
}

SEVERITY_LABEL = {
    "HIGH":   "⚠ HIGH",
    "MEDIUM": "△ MEDIUM",
    "LOW":    "○ LOW",
    "PASS":   "✓ PASS",
}


# ---------------------------------------------------------------------------
# Style configuration
#EDIT:04/06, 14:55: Put old configure_styles back
# ---------------------------------------------------------------------------
def configure_styles(doc: Document) -> None:
    # Page: A4 with 0.5-inch margins
    for section in doc.sections:
        section.page_width   = Inches(8.27)
        section.page_height  = Inches(11.69)
        section.left_margin  = section.right_margin  = Inches(0.5)
        section.top_margin   = section.bottom_margin = Inches(0.5)
        section.footer_distance = Inches(0.4)

    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.font.color.rgb = DARK_GREY

    h1 = styles["Heading 1"]
    h1.font.name = "Arial"
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after  = Pt(4)

    h2 = styles["Heading 2"]
    h2.font.name = "Arial"
    h2.font.size = Pt(11)
    h2.font.bold = True
    h2.font.color.rgb = NAVY
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after  = Pt(3)

    h3 = styles["Heading 3"]
    h3.font.name = "Arial"
    h3.font.size = Pt(10)
    h3.font.bold = True
    h3.font.color.rgb = DARK_GREY
    h3.paragraph_format.space_before = Pt(6)
    h3.paragraph_format.space_after  = Pt(2)

# ---------------------------------------------------------------------------
# Low-level paragraph / run helpers
# ---------------------------------------------------------------------------

def add_page_break(doc: Document) -> None:
    p   = doc.add_paragraph()
    run = p.add_run()
    br  = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)

#04/06 edit, 15:20, added back the previous add_heading
def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)

# 04/06 edit, 15:20, added back the previous
def add_body(doc: Document, text: str, italic: bool = False,
             bold: bool = False, colour: RGBColor = None) -> None:
    p   = doc.add_paragraph(style="Normal")
    run = p.add_run(text)
    run.italic = italic
    run.bold   = bold
    if colour:
        run.font.color.rgb = colour
    p.paragraph_format.space_after = Pt(4)

# 04/06 edit, 15:22, added back the previous add_bullet
def add_bullet(doc: Document, text: str, level: int = 0,
               colour: RGBColor = None, bold_prefix: str = None) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        if colour:
            r.font.color.rgb = colour
    r2 = p.add_run(text)
    if colour:
        r2.font.color.rgb = colour

def add_labelled_bullet(doc: Document, label: str, value: str,
                        colour: RGBColor = None) -> None:
    p  = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    if colour:
        r1.font.color.rgb = colour
    p.add_run(value)


# 04/06 edit, 15:23, added back the previous add_placeholder
def add_placeholder(doc: Document, section_hint: str) -> None:
    """Insert a greyed italic placeholder when no slide content is available."""
    p   = doc.add_paragraph(style="Normal")
    run = p.add_run(
        f"[{section_hint} — content extracted from pitch deck. "
        "Review and edit as required.]"
    )
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    p.paragraph_format.space_after = Pt(6)


def add_gap_note(doc: Document, gap_text: str) -> None:
    """
    Insert a clearly visible gap / missing-information note.
    Used to flag where the pitch deck did not provide expected DD content.
    """
    p   = doc.add_paragraph(style="Normal")
    run = p.add_run(f"⚠ Information gap: {gap_text}")
    run.italic = True
    run.bold   = True
    run.font.color.rgb = AMBER
    p.paragraph_format.space_after = Pt(4)


# 04/06 edit, 15:25, added back previous add_mixed_section
def add_mixed_section(doc: Document, text: str) -> None:
    """
    Render LLM output that contains a mix of prose paragraphs and bullet lines.

    Rules applied:
    - Lines starting with "•" or "-" are rendered as list bullets.
    - Lines that are sub-headers (end with ":") are rendered as bold body text.
    - All other non-empty lines are rendered as body paragraphs.
    - Consecutive blank lines are collapsed to a single spacer paragraph.

    This replaces the previous pattern of calling add_body() for every line,
    which lost bullet structure when the LLM returned bullet-formatted output.
    """
    if not text or text.strip() == "[LLM unavailable — review and complete this section manually.]":
        add_placeholder(doc, "Review and complete this section manually")
        return

    prev_blank = False
    for raw_line in text.splitlines():
        line = raw_line.strip()

        if not line:
            if not prev_blank:
                # Single spacer — keep spacing tight
                sp = doc.add_paragraph(style="Normal")
                sp.paragraph_format.space_after = Pt(2)
            prev_blank = True
            continue

        prev_blank = False

        # Bullet lines
        if line.startswith("• ") or line.startswith("- "):
            content = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(content)
            r.font.size = Pt(10)
            r.font.color.rgb = DARK_GREY
            p.paragraph_format.space_after = Pt(3)

        # Sub-header lines (e.g. "Gaps / Concerns:" or "Red Flags:")
        elif line.endswith(":") and len(line) < 60:
            p   = doc.add_paragraph(style="Normal")
            run = p.add_run(line)
            run.bold           = True
            run.font.size      = Pt(10)
            run.font.color.rgb = NAVY
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)

        # Regular prose paragraph
        else:
            add_body(doc, line)

#edit: add_questions_list() for LLM output that is expected to be a numbered list of questions, with an optional intro sentence.
def add_questions_list(doc: Document, text: str) -> None:
    """
    Render LLM output that is structured as:
      One introductory sentence.
      1. Question one
      2. Question two
      ...

    The intro sentence is written as a body paragraph.
    Numbered lines are indented as numbered list items.
    Falls back to add_mixed_section if no numbered lines are detected.
    """
    if not text or text.strip() == _FALLBACK_NOTE_SENTINEL:
        add_placeholder(doc, "Review and complete this section manually")
        return

    lines = [l.strip() for l in text.splitlines() if l.strip()]
    has_numbered = any(l[:3].rstrip(". ").isdigit() for l in lines)
    if not has_numbered:
        add_mixed_section(doc, text)
        return

    for line in lines:
        # Detect "N." or "N)" prefix
        stripped = line.lstrip("0123456789")
        if stripped and stripped[0] in ".)" and len(line) != len(stripped):
            question_text = stripped[1:].strip()
            p = doc.add_paragraph(style="List Number")
            r = p.add_run(question_text)
            r.font.size      = Pt(10)
            r.font.color.rgb = DARK_GREY
            p.paragraph_format.space_after = Pt(3)
        else:
            # Introductory sentence or sub-header
            add_body(doc, line)

# ---------------------------------------------------------------------------
# Table helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_colour: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd   = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tc_pr.append(shd)


def _set_col_width(cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w  = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"),    str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def add_findings_table(doc: Document, findings: list, max_rows: int = 50) -> None:
    """Render a colour-coded anomaly findings table."""
    if not findings:
        add_body(doc, "No findings in this category.", italic=True)
        return

    headers    = ["ID", "Category", "Severity", "Label", "Evidence", "Slide"]
    col_widths = [0.7,   1.1,        0.85,       1.6,     4.0,        0.6]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Header row
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = table.rows[0].cells[i]
        _set_cell_bg(cell, LIGHT_BLUE_FILL)
        _set_col_width(cell, w)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = NAVY

    for idx, f in enumerate(findings[:max_rows]):
        row = table.add_row()
        sev = getattr(f, "severity", "")
        row_fill = {"HIGH": "FFE0E0", "MEDIUM": "FFF3CD", "LOW": "E8F5E9"}.get(sev, "FFFFFF")

        vals = [
            getattr(f, "anomaly_id",      ""),
            getattr(f, "category",        ""),
            SEVERITY_LABEL.get(sev, sev),
            getattr(f, "label",           ""),
            (getattr(f, "evidence", "") or "")[:200],
            str(getattr(f, "slide_reference", "") or ""),
        ]

        for i, (v, w) in enumerate(zip(vals, col_widths)):
            cell = row.cells[i]
            fill = ALT_ROW_FILL if (idx % 2 == 1 and sev == "") else row_fill
            _set_cell_bg(cell, fill)
            _set_col_width(cell, w)
            run = cell.paragraphs[0].add_run(str(v))
            run.font.size = Pt(8.5)
            if i == 2:  # severity column — coloured text
                run.font.color.rgb = SEVERITY_COLOUR.get(sev, DARK_GREY)
                run.bold = True

    doc.add_paragraph()


def add_questions_table(doc: Document, questions_by_cat: dict) -> None:
    """Render investor questions grouped by category."""
    if not questions_by_cat:
        add_body(doc, "No questions generated.", italic=True)
        return
    for cat, qs in questions_by_cat.items():
        add_heading(doc, cat, level=3)
        for q in qs:
            add_bullet(doc, str(q))


# ADD: 24/05
def extract_company_name(result) -> str:
    """
    Extract company name from deck content in priority order:
    1. result.deck.company_name  — if the pipeline sets it explicitly
    2. Scan ALL slides (any section label) for the first credible name line
    3. Source filename stripped of extension and underscores
    """
    # ── Priority 1: pipeline-provided field ──────────────────────────────
    name = getattr(getattr(result, "deck", object()), "company_name", "") or ""
    if name.strip():
        return name.strip().title()

    # ── Priority 2: scan slide content ───────────────────────────────────
    # Words that appear on cover/disclaimer slides but are NOT company names
    _SKIP_STARTS = (
        "private", "confidential", "disclaimer", "the material",
        "this presentation", "prepared by", "for limited",
        "early investment", "teaser", "seed", "series",
        "preliminary", "investment", "venture", "startupscale",
        "section", "contents", "table of", "agenda",
    )
    _SKIP_EXACT = {
        "ltd", "llc", "inc", "fze", "fzc", "plc", "limited",
        "deck", "slide", "page", "overview", "summary",
    }

    try:
        slides = result.deck.slides
        # Check COVER slides first, then fall through to all slides
        cover_first = sorted(
            slides,
            key=lambda s: 0 if (getattr(s, "detected_section", "") or "").upper()
                              in ("COVER", "ABOUT", "INTRO", "OVERVIEW") else 1
        )

        for slide in cover_first:
            txt = (getattr(slide, "cleaned_text", "")
                   or getattr(slide, "raw_text", "") or "")
            for line in txt.splitlines():
                line = line.strip()
                lower = line.lower()

                # Must be between 2 and 60 characters
                if not (2 <= len(line) <= 60):
                    continue

                # Skip pure numbers or slide references
                if line.replace(" ", "").isdigit():
                    continue

                # Skip lines that start with known non-name words
                if any(lower.startswith(s) for s in _SKIP_STARTS):
                    continue

                # Skip lines that ARE a known non-name word
                if lower in _SKIP_EXACT:
                    continue

                # Skip lines that look like sentences (contain a verb signal)
                if len(line.split()) > 6:
                    continue

                # Skip lines that are ALL punctuation or symbols
                if all(not c.isalpha() for c in line):
                    continue

                # Skip email addresses and URLs
                if "@" in line or "www." in lower or "http" in lower:
                    continue

                # This line is a credible company name candidate
                return line

    except AttributeError:
        pass

    # ── Priority 3: filename fallback ─────────────────────────────────────
    source = getattr(getattr(result, "deck", object()),
                     "source_filename", "Unknown")
    return (
        source.replace(".pdf", "")
              .replace(".pptx", "")
              .replace("_", " ")
              .replace("-", " ")
              .strip()
              .title()
    )

# ---------------------------------------------------------------------------
# Edit: 04/06, 14:59, put old title_block back
# Title block — replaces cover page when no full cover is needed
# ---------------------------------------------------------------------------

def title_block(doc: Document, company_name: str, report_type: str,
                source_filename: str, prepared_for: str = "") -> None:
    """
    Render the report title block at the top of page 1.
    
    section index lines / meta line. No page break, no cover image.
    """
    def _p(space_after_pt):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(space_after_pt)
        return p

    # Line 1: report label — 14pt teal bold
    #p1 = _p(6)
    #r1 = p1.add_run("CONSOLIDATED INVESTMENT MEMO")
    #r1.bold = True; r1.font.size = Pt(14)
    #r1.font.name = "Arial"; r1.font.color.rgb = NAVY

    # Line 2: company name — 18pt black bold  ← specified requirement
    p2 = _p(6)
    r2 = p2.add_run(company_name.upper())
    r2.bold = True; r2.font.size = Pt(18)
    r2.font.name = "Arial"; r2.font.color.rgb = BLACK

    # Line 3: report type subtitle — 12pt grey
    p3 = _p(6)
    r3 = p3.add_run(report_type)
    r3.font.size = Pt(12); r3.font.name = "Arial"
    r3.font.color.rgb = DARK_GREY

    # Lines 4-5: section index — 10pt teal
    for line in [
        "Section A: All Investment Rationale Points (by Source Document)",
        "Section B: All Due Diligence Questions (Prioritised by Severity)",
    ]:
        p = _p(3)
        r = p.add_run(line)
        r.font.size = Pt(10); r.font.name = "Arial"
        r.font.color.rgb = NAVY

    # Line 6: meta — 9pt grey, larger gap below
    p6 = _p(32)
    date_str = datetime.now().strftime("%d %B %Y").lstrip("0")
    meta = f"{date_str}  |   Source: {source_filename}"
    if prepared_for:
        meta += f"  |  Prepared for {prepared_for}"
    r6 = p6.add_run(meta)
    r6.font.size = Pt(9); r6.font.name = "Arial"
    r6.font.color.rgb = LIGHT_GREY


# ---------------------------------------------------------------------------
# Card helpers — investment rationale and DD question cards
# ADDED: 24/05
# ---------------------------------------------------------------------------

def add_card(doc: Document, number: int, category: str,
             source: str, body_text: str,
             severity: str = None) -> None:
    """
    Render a single investment rationale or DD question card matching
    light header cell, full-width merged body row.

    severity: None (rationale), 'CRITICAL', 'HIGH', 'MEDIUM'
    """
    severity_badges  = {"CRITICAL": "  🚨 CRITICAL", "HIGH": "  ⚠️ HIGH", "MEDIUM": "  ℹ️ MEDIUM"}
    category_colours = {"CRITICAL": RED, "HIGH": AMBER, "MEDIUM": GREEN}

    cat_colour = category_colours.get(severity, NAVY) if severity else NAVY
    badge      = severity_badges.get(severity, "")    if severity else ""

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"

    # Row 0 — number cell
    cell_num = table.rows[0].cells[0]
    _set_cell_bg(cell_num, CARD_NUMBER_FILL)
    _set_col_width(cell_num, _CARD_NUM_COL / 1440)
    p_num = cell_num.paragraphs[0]
    p_num.paragraph_format.space_before = Pt(4)
    p_num.paragraph_format.space_after  = Pt(4)
    r_num = p_num.add_run(str(number))
    r_num.bold = True; r_num.font.size = Pt(14)
    r_num.font.name = "Arial"
    r_num.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Row 0 — header cell
    cell_hdr = table.rows[0].cells[1]
    _set_cell_bg(cell_hdr, CARD_HEADER_FILL)
    _set_col_width(cell_hdr, _CARD_BODY_COL / 1440)

    p_cat = cell_hdr.paragraphs[0]
    p_cat.paragraph_format.space_before = Pt(3)
    p_cat.paragraph_format.space_after  = Pt(1)
    r_cat = p_cat.add_run(f"{category.upper()}{badge}")
    r_cat.bold = True; r_cat.font.size = Pt(11)
    r_cat.font.name = "Arial"; r_cat.font.color.rgb = cat_colour

    p_src = cell_hdr.add_paragraph()
    p_src.paragraph_format.space_before = Pt(0)
    p_src.paragraph_format.space_after  = Pt(3)
    r_src_l = p_src.add_run("Source Document: ")
    r_src_l.font.size = Pt(9); r_src_l.font.name = "Arial"
    r_src_l.font.color.rgb = DARK_GREY
    r_src_v = p_src.add_run(source)
    r_src_v.font.size = Pt(9); r_src_v.font.name = "Arial"
    r_src_v.font.color.rgb = DARK_GREY

    # Row 1 — full-width merged body
    table.rows[1].cells[0].merge(table.rows[1].cells[1])
    cell_body = table.rows[1].cells[0]
    _set_col_width(cell_body, (_CARD_NUM_COL + _CARD_BODY_COL) / 1440)
    p_body = cell_body.paragraphs[0]
    p_body.paragraph_format.space_before = Pt(5)
    p_body.paragraph_format.space_after  = Pt(5)
    r_body = p_body.add_run(body_text)
    r_body.font.size = Pt(10); r_body.font.name = "Arial"

    # Gap after card
    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(0)
    gap.paragraph_format.space_after  = Pt(4)


def add_severity_header(doc: Document, severity: str, count: int,
                        subtitle: str) -> None:
    """
    Render a full-width coloured severity section divider
    (e.g. 🚨 CRITICAL — 5 questions).
    """
    fill_map  = {"CRITICAL": CRITICAL_FILL, "HIGH": HIGH_FILL, "MEDIUM": MEDIUM_FILL}
    badge_map = {"CRITICAL": "🚨 CRITICAL",  "HIGH": "⚠️ HIGH",  "MEDIUM": "ℹ️ MEDIUM"}

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    _set_cell_bg(cell, fill_map.get(severity, CRITICAL_FILL))
    _set_col_width(cell, (_CARD_NUM_COL + _CARD_BODY_COL) / 1440)

    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(5)
    p1.paragraph_format.space_after  = Pt(2)
    r1 = p1.add_run(f"{badge_map.get(severity, severity)} ({count} questions)")
    r1.bold = True; r1.font.size = Pt(12)
    r1.font.name = "Arial"
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(5)
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(9); r2.font.name = "Arial"
    r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ---------------------------------------------------------------------------
# Cover page EDIT:10/04
# ---------------------------------------------------------------------------
def cover_page(doc: Document, title: str, company_name: str,
               source_filename: str,
               hero_image_path: str = None,
               logo_image_path: str = None) -> None:
    """
    Build the StartUpScale360 cover page matching the reference screenshot:

      1. Logo — centred at top (image if logo_image_path given, else styled text)
      2. Hero image — full-width grey photo (image if hero_image_path given, else shaded box)
      3. Cyan banner — report title in white bold caps (overlaid lower-left of image zone)
      4. Company name — in brand cyan, centred, below the banner
      5. Date — small grey text, centred, at the very bottom
      6. Page break

    Parameters
    ----------
    title            : Report type, e.g. "Preliminary Due Diligence Report"
    company_name     : The startup being reviewed, e.g. "MIRA" — shown in cyan below banner
    source_filename  : Original deck filename (used internally, not shown on cover)
    hero_image_path  : Optional path to a cover photo (JPEG/PNG). If omitted, a grey
                       placeholder box is used.
    logo_image_path  : Optional path to the StartUpScale360 logo PNG. If omitted, a
                       two-tone text logo is rendered.
    """

    # ── 1. Logo ────────────────────────────────────────────────────────────
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(1) # edit 05/06, 13:21, reduced space above logo
    p_logo.paragraph_format.space_after  = Pt(1)

    if logo_image_path and os.path.isfile(logo_image_path):
        p_logo.add_run().add_picture(logo_image_path, width=Inches(2.2))
    else:
        # Two-tone typographic logo: START (black) UP (cyan) SCALE (black) 360 (black)
        def _lr(text, colour, size=15, italic=False):
            r = p_logo.add_run(text)
            r.bold = True
            r.italic = italic
            r.font.name = "Arial"
            r.font.size = Pt(size)
            r.font.color.rgb = colour

        _lr("START",  NAVY)
        _lr("UP",     CYAN)
        _lr("SCALE",  NAVY)
        # Arc glyph approximated with a rotated arrow; swap for an image logo when available
        _lr(" \u21BA", CYAN, size=12)   # ↺ character
        _lr("360",    NAVY)

    # ── 2. Hero image ──────────────────────────────────────────────────────
    if hero_image_path and os.path.isfile(hero_image_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(0)
        p_img.paragraph_format.space_after  = Pt(0)
        p_img.add_run().add_picture(hero_image_path, width=Inches(5.8))
    else:
        # Grey shaded placeholder table (mimics the dark photo)
        tbl = doc.add_table(rows=1, cols=1)
        cell = tbl.rows[0].cells[0]
        _set_cell_bg(cell, "404040")          # dark grey fill
        _set_col_width(cell, 5.8)
        # Set row height ~3.2 inches
        tr_pr = tbl.rows[0]._tr.get_or_add_trPr()
        tr_h  = OxmlElement("w:trHeight")
        tr_h.set(qn("w:val"),   str(int(3.2 * 1440)))
        tr_h.set(qn("w:hRule"), "exact")
        tr_pr.append(tr_h)
        # Italic placeholder text centred in the box
        ph = cell.paragraphs[0]
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ph.paragraph_format.space_before = Pt(60)
        r_ph = ph.add_run("[Add cover image here]")
        r_ph.italic = True
        r_ph.font.size = Pt(10)
        r_ph.font.color.rgb = WHITE

    # ── 3. Cyan banner — report title ──────────────────────────────────────
    # Rendered as a single-cell table with cyan background so it spans the
    # left portion of the page like the banner in the reference screenshot.
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    banner_tbl = doc.add_table(rows=1, cols=1)
    b_cell = banner_tbl.rows[0].cells[0]
    _set_cell_bg(b_cell, "00B0F0")            # brand cyan fill
    _set_col_width(b_cell, 3.6)               # left ~60% of page width

    b_tr_pr = banner_tbl.rows[0]._tr.get_or_add_trPr()
    b_tr_h  = OxmlElement("w:trHeight")
    b_tr_h.set(qn("w:val"),   str(int(0.75 * 1440)))
    b_tr_h.set(qn("w:hRule"), "atLeast")
    b_tr_pr.append(b_tr_h)

    b_p = b_cell.paragraphs[0]
    b_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    b_p.paragraph_format.space_before = Pt(6)
    b_p.paragraph_format.space_after  = Pt(6)

    # Write title as upper-case lines (split on newline if caller passes one)
    for line in title.upper().split("\n"):
        r_t = b_p.add_run(line + "\n")
        r_t.bold           = True
        r_t.font.name      = "Arial"
        r_t.font.size      = Pt(15)
        r_t.font.color.rgb = WHITE

    # ── 4. Company name in cyan ────────────────────────────────────────────
    p_co = doc.add_paragraph()
    p_co.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_co.paragraph_format.space_before = Pt(10)
    p_co.paragraph_format.space_after  = Pt(20)
    r_co = p_co.add_run(company_name.upper())
    r_co.bold           = True
    r_co.font.name      = "Arial"
    r_co.font.size      = Pt(18)
    r_co.font.color.rgb = CYAN

    # ── 5. Date ────────────────────────────────────────────────────────────
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(4)
    p_date.paragraph_format.space_after  = Pt(8)
    r_date = p_date.add_run(datetime.now().strftime("%d %b %Y").lstrip("0"))
    r_date.font.name      = "Arial"
    r_date.font.size      = Pt(9)
    r_date.font.color.rgb = LIGHT_GREY

    # ── 6. Page break ──────────────────────────────────────────────────────
    add_page_break(doc)

# ---------------------------------------------------------------------------
# Inner-page header: logo top-right
# ---------------------------------------------------------------------------

def _build_logo_paragraph(paragraph, logo_image_path: str = None,
                           logo_width_inches: float = 1.4) -> None:
    """
    Add the StartUpScale360 logo into `paragraph`, right-aligned.
    Uses an image if logo_image_path is provided and exists, otherwise
    renders the two-tone text logo as a fallback.
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after  = Pt(0) #edit: 04/06, 15:52, reduced space after to tighten header layout, 2 to 0

    if logo_image_path and os.path.isfile(logo_image_path):
        paragraph.add_run().add_picture(logo_image_path, width=Inches(logo_width_inches))
    else:
        def _lr(text, colour, size=10):
            r = paragraph.add_run(text)
            r.bold = True
            r.font.name = "Arial"
            r.font.size = Pt(size)
            r.font.color.rgb = colour

        _lr("START",  NAVY)
        _lr("UP",     CYAN)
        _lr("SCALE",  NAVY)
        _lr(" \u21BA", CYAN, size=8)
        _lr("360",    NAVY)


def setup_page_header(section, logo_image_path: str = None) -> None:
    """
    Set the inner-page header for `section`:
      - Logo (image or text) right-aligned
      - Thin cyan rule beneath

    Call once per section after setting different_first_page_header_footer = True
    so the cover page is unaffected.

    Usage in framework_a.py / framework_b.py:
        section = doc.sections[0]
        section.different_first_page_header_footer = True
        setup_page_header(section, logo_image_path="path/to/logo.png")
        setup_page_footer(section)
    """
    header = section.header
    header.is_linked_to_previous = False

    # Clear any default content
    for p in header.paragraphs:
        for r in p.runs:
            r.text = ""

    # Logo paragraph
    logo_p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    _build_logo_paragraph(logo_p, logo_image_path)

    # Thin cyan rule beneath the logo
    rule_p = header.add_paragraph()
    rule_p.paragraph_format.space_before = Pt(0)
    rule_p.paragraph_format.space_after  = Pt(0)
    pPr  = rule_p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")          # ¾ pt line
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "00B0F0")     # brand cyan
    pBdr.append(bot)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# EDIT: 10/04 Inner-page footer: address / contact / copyright
# ---------------------------------------------------------------------------

def setup_page_footer(section) -> None:
    """
    Set the inner-page footer for `section`:

      Line 1 │ StartUpScale360 FZE   (bold, dark grey)
      Line 2 │ Tel: …  |  Address: …  |  Website: …  |  Email: …
      ─────── thin grey rule ────────────────────────────────────────
      Line 3 │ © copyright disclaimer  (tiny italic grey)

    Edit COMPANY_TEL / COMPANY_ADDRESS / COMPANY_WEBSITE / COMPANY_EMAIL /
    COPYRIGHT_LINE at the top of this file to update all documents at once.
    """
    footer = section.footer
    footer.is_linked_to_previous = False

    # Clear default paragraphs
    for p in footer.paragraphs:
        for r in p.runs:
            r.text = ""

    def _add_footer_para(space_before=1, space_after=2):
        p = footer.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        return p

    # ── Line 1: company name ──
    p1 = footer.paragraphs[0] if footer.paragraphs else _add_footer_para()
    p1.clear()
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after  = Pt(1)
    r1 = p1.add_run(COMPANY_NAME)
    r1.bold           = True
    r1.font.name      = "Arial"
    r1.font.size      = Pt(8.5)
    r1.font.color.rgb = NAVY

    # ── Line 2: contact details ──
    p2 = _add_footer_para(space_before=0, space_after=3)

    def _contact(label, value, sep=" | "):
        rb = p2.add_run(f"{label}: ")
        rb.bold           = True
        rb.font.name      = "Arial"
        rb.font.size      = Pt(7.5)
        rb.font.color.rgb = DARK_GREY
        rv = p2.add_run(f"{value}{sep}")
        rv.font.name      = "Arial"
        rv.font.size      = Pt(7.5)
        rv.font.color.rgb = LIGHT_GREY

    _contact("Tel",     COMPANY_TEL)
    _contact("Address", COMPANY_ADDRESS)
    _contact("Website", COMPANY_WEBSITE)
    _contact("Email",   COMPANY_EMAIL, sep="")   # no trailing separator on last item

    # ── Thin grey rule ──
    p_rule = _add_footer_para(space_before=2, space_after=2)
    pPr    = p_rule._p.get_or_add_pPr()
    pBdr   = OxmlElement("w:pBdr")
    top    = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "AAAAAA")
    pBdr.append(top)
    pPr.append(pBdr)

    # ── Line 3: copyright ──
    p3  = _add_footer_para(space_before=0, space_after=0)
    r3  = p3.add_run(COPYRIGHT_LINE)
    r3.italic         = True
    r3.font.name      = "Arial"
    r3.font.size      = Pt(6.5)
    r3.font.color.rgb = LIGHT_GREY

# ---------------------------------------------------------------------------
# Slide content extraction
# ---------------------------------------------------------------------------

def _normalise(s: str) -> str:
    """Lowercase and replace underscores with spaces for fuzzy section matching."""
    return (s or "").lower().replace("_", " ")


def slides_for_section(result, section_name: str) -> list:
    """
    Return slides whose detected_section matches section_name.

    Normalises both sides to lowercase-with-spaces before comparing, so
    pipeline labels like 'CAP_TABLE' match search terms like 'cap table',
    and 'SOLUTION' matches 'solution'. Slides classified as UNKNOWN are
    always excluded.
    """
    try:
        slides = result.deck.slides
    except AttributeError:
        return []

    needle = _normalise(section_name)
    matched = []
    for s in slides:
        raw  = getattr(s, "detected_section", "") or ""
        norm = _normalise(raw)
        if norm == "unknown":
            continue
        if needle in norm or norm in needle:
            matched.append(s)
    return matched


def emit_slide_content(doc: Document, result, section_name: str,
                       placeholder_hint: str) -> None:
    """Write slide text as body paragraphs, or a placeholder if nothing found."""
    slides = slides_for_section(result, section_name)
    parts  = []
    for s in slides:
        txt = getattr(s, "cleaned_text", "") or getattr(s, "raw_text", "") or ""
        if txt.strip():
            parts.append(txt.strip())

    if parts:
        for para in "\n\n".join(parts).split("\n"):
            if para.strip():
                add_body(doc, para.strip())
    else:
        add_placeholder(doc, placeholder_hint)


def safe_call(result, method: str, *args, default=None):
    """Call result.method(*args) safely, returning default on any error."""
    try:
        fn = getattr(result, method)
        return fn(*args) if args else fn()
    except Exception:
        return default if default is not None else []